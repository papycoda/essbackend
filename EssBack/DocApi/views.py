from io import BytesIO
from rest_framework import viewsets, status
from rest_framework.views import APIView
from rest_framework.decorators import action
from rest_framework.response import Response
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from django.views.generic import View
from rest_framework.permissions import IsAuthenticated, IsAdminUser
from .models import ESGProject, UserTask
from .serializers import ESGProjectSerializer, UserTaskSerializer, UserCreateSerializer
from rest_framework.decorators import permission_classes, api_view


class ESGProjectViewSet(viewsets.ModelViewSet):
    queryset = ESGProject.objects.all()
    serializer_class = ESGProjectSerializer

class UserTaskViewSet(viewsets.ModelViewSet):
    queryset = UserTask.objects.all()
    serializer_class = UserTaskSerializer

@api_view(['POST'])
@permission_classes([IsAuthenticated, IsAdminUser])
def create_user(request):
    if request.method == 'POST':
        serializer = UserCreateSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class GenerateReportView(View):
    def get(self, request, format, project_id):
        project = get_object_or_404(ESGProject, pk=project_id)

        if format == 'docx':
            response = self.generate_docx_report(project)
            filename = f'{project.company_name}_report.docx'
        elif format == 'pdf':
            response = self.generate_pdf_report(project)
            filename = f'{project.company_name}_report.pdf'
        else:
            return HttpResponse('Unsupported format', status=400)

        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    def generate_docx_report(self, project):
        doc = Document()
        doc.add_heading(f'{project.company_name} Project Report', level=1)
        self.add_project_info(doc, project)
        doc.add_paragraph('Summary of the project goes here.')
        buffer = BytesIO()
        doc.save(buffer)
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="report.docx"'
        return response

    def generate_pdf_report(self, project):
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{project.company_name}_report.pdf'

        c = canvas.Canvas(response, pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(100, 750, f'{project.company_name} Project Report')
        self.add_project_info_pdf(c, project)
        c.drawString(100, 600, 'Summary of the project goes here')
        c.showPage()
        c.save()
        return response

    def add_project_info(self, doc, project):
        doc.add_paragraph(f'Company Name: {project.company_name}')
        doc.add_paragraph(f'Sector: {project.sector}')
        doc.add_paragraph(f'Location: {project.location}')
        doc.add_paragraph(f'Owner: {project.owners_name}')
        doc.add_paragraph(f'Start Date: {project.start_date}')

    def add_project_info_pdf(self, c, project):
        c.drawString(100, 700, f'Company Name: {project.company_name}')
        c.drawString(100, 680, f'Sector: {project.sector}')
        c.drawString(100, 660, f'Location: {project.location}')
        c.drawString(100, 640, f'Owner: {project.owners_name}')
        c.drawString(100, 620, f'Start Date: {project.start_date}')